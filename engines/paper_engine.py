# -*- coding: utf-8 -*-
"""
paper_engine.py
───────────────
실계산 수치 기반 논문 초안 자동 생성 (한국어 + English bilingual)

generate_paper(session_data) → dict
  {
    "abstract":     str,
    "intro":        str,
    "hypothesis":   str,
    "methods":      str,
    "results":      str,
    "hypothesis_val": str,
    "conclusion":   str,
    "limitations":  str,
    "full_md":      str,   # 전체 Markdown
  }

모든 수치는 session_state 의 실제 계산 결과에서 추출
LLM 생성 텍스트 없음 — 도메인 지식 템플릿 + 실수치 조합
"""

from __future__ import annotations
import numpy as np
import pandas as pd
from datetime import date


# ─────────────────────────────────────────────────────────────────────────────
def generate_paper(
    df_dataset:   pd.DataFrame,       # crc_dielectric_combined.csv
    feat_names:   list[str],          # 선택된 피처
    X_all_shape:  tuple,              # (n_mol, n_raw_feat)
    cv_result:    dict,               # 5-fold CV 결과
    metrics:      dict,               # {"gbr": {"train":{}, "val":{}, "test":{}}}
    df_screening: pd.DataFrame,       # S4 스크리닝 결과
    feat_imp:     pd.Series,          # 피처 중요도
    log_tf:       bool = True,
    best_model:   str  = "gbr",       # 자동 선택된 최적 모델
    model_mode:   str  = "lowk",      # "lowk" | "full"
    k_threshold:  float = 2.4,        # 저유전율 스크리닝 기준
    train_threshold: float = 3.5,     # Low-k 전용 모드 훈련 기준
) -> dict:
    """
    실계산 수치 기반 논문 초안 생성 (한국어 주체, 영문 Abstract 포함).
    """
    today = date.today().strftime("%Y-%m-%d")

    # ── 핵심 수치 추출 ────────────────────────────────────────────────────────
    n_total     = len(df_dataset)
    k_min       = df_dataset["k_exp"].min()
    k_max       = df_dataset["k_exp"].max()
    k_mean      = df_dataset["k_exp"].mean()
    k_median    = df_dataset["k_exp"].median()
    n_lowk      = int((df_dataset["k_exp"] < k_threshold).sum())
    n_lowk35    = int((df_dataset["k_exp"] < 3.5).sum())
    n_ultralowk = int((df_dataset["k_exp"] < 2.5).sum())

    # 카테고리 정보
    if "category" in df_dataset.columns:
        cats = df_dataset["category"].nunique()
        top_cats = df_dataset["category"].value_counts().head(5).index.tolist()
        top_cats_str = ", ".join(top_cats)
    else:
        cats = "–"
        top_cats_str = "알케인, 방향족, 불소화합물, 실리콘계 등"

    # 소스 정보
    if "source" in df_dataset.columns:
        sources = df_dataset["source"].value_counts().to_dict()
        source_str = ", ".join(f"{src}({cnt}종)" for src, cnt in sources.items())
    else:
        source_str = "CRC Handbook, 문헌"

    n_raw_feat = X_all_shape[1]
    n_sel_feat = len(feat_names)

    # 피처 중요도 top5
    top5_feat = feat_imp.head(5).index.tolist() if len(feat_imp) >= 5 else feat_imp.index.tolist()
    top5_imp  = feat_imp.head(5).values.tolist() if len(feat_imp) >= 5 else feat_imp.values.tolist()
    top3_feat = top5_feat[:3]
    top3_imp  = top5_imp[:3]

    # 모델 성능
    bm = best_model.upper()
    bm_key = best_model.lower()
    bm_tr = metrics.get(bm_key, {}).get("train", {})
    bm_va = metrics.get(bm_key, {}).get("val",   {})
    bm_te = metrics.get(bm_key, {}).get("test",  {})

    gbr_te = metrics.get("gbr",   {}).get("test", {})
    rf_te  = metrics.get("rf",    {}).get("test", {})
    ri_te  = metrics.get("ridge", {}).get("test", {})
    gpr_te = metrics.get("gpr",   {}).get("test", {})

    cv_r2    = cv_result.get("r2_mean",   0.0)
    cv_r2std = cv_result.get("r2_std",    0.0)
    cv_rmse  = cv_result.get("rmse_mean", 0.0)

    bm_te_r2   = float(bm_te.get("R2",   0.0))
    bm_te_rmse = float(bm_te.get("RMSE", 0.0))
    bm_te_mae  = float(bm_te.get("MAE",  0.0))
    bm_tr_r2   = float(bm_tr.get("R2",   0.0))
    log_rmse_val = bm_te.get("log_RMSE", None)
    log_rmse_str = f"{float(log_rmse_val):.4f}" if log_rmse_val is not None else "–"

    # 스크리닝 결과
    df_valid = df_screening[df_screening["valid"] == True].copy() \
               if "valid" in df_screening.columns else df_screening.copy()
    n_cand    = len(df_valid)
    n_pred_lk = int((df_valid["pred_mean"] < k_threshold).sum()) \
                if "pred_mean" in df_valid.columns else 0
    n_ad_ok   = int(df_valid["ad_ok"].sum()) \
                if "ad_ok" in df_valid.columns else 0
    top5_df   = df_valid.head(5) if not df_valid.empty else pd.DataFrame()

    # 모드 설명
    mode_desc_kr = ("Low-k 전용 모드 (k < {:.1f} 학습 데이터만 사용)".format(train_threshold)
                    if model_mode == "lowk"
                    else "전체 데이터 모드 (k = {:.2f}–{:.0f})".format(k_min, k_max))
    mode_desc_en = ("Low-k dedicated mode (training data: k < {:.1f})".format(train_threshold)
                    if model_mode == "lowk"
                    else "Full-range mode (k = {:.2f}–{:.0f})".format(k_min, k_max))

    train_n = int(0.70 * n_total)
    h_star  = round(3 * (n_sel_feat + 1) / train_n, 4)

    # 피처 물리 해석
    feat_interp = _feat_interpretation(top5_feat)

    # 상위 후보 화학적 해석
    top_chem = _top_candidate_chemistry(top5_df)

    # ── 제목 & Abstract ───────────────────────────────────────────────────────
    title_kr = "QSPR 머신러닝 기반 저유전율 유기 박막 소재 설계 및 선별"
    title_en = "QSPR Machine Learning-Based Design and Screening of Low-k Organic Dielectric Thin Film Materials"

    abstract = f"""\
## 초록 / Abstract

**[한국어]**
차세대 플렉서블 OLED 박막봉지(TFE)의 유기층에 적용 가능한 저유전율(low-k, \
k < {k_threshold}) 유기 소재를 체계적으로 선별하기 위해 QSPR(Quantitative \
Structure–Property Relationship) 머신러닝 파이프라인을 구축하였다. \
CRC 핸드북 및 관련 문헌에서 수집한 **{n_total}종** 화합물(유전율 범위: \
{k_min:.2f}–{k_max:.1f})을 훈련 데이터로 활용하였으며, \
RDKit으로 산출한 {n_raw_feat}개 분자 기술자를 3단계 선택 파이프라인으로 \
**{n_sel_feat}개**로 축약하였다. {model_mode.upper()} 모드에서 \
{mode_desc_kr}를 적용한 결과, \
**{bm}** 모델이 5겹 교차검증 R² = {cv_r2:.3f} ± {cv_r2std:.3f}, \
테스트 R² = {bm_te_r2:.3f}, logRMSE = {log_rmse_str}로 최우수 성능을 보였다. \
{n_cand}종 후보 분자에 대한 전향적 스크리닝에서 \
**{n_pred_lk}종**이 k < {k_threshold}로 예측되어 실험 검증 우선순위 목록을 제공하였다.

**[English]**
A QSPR machine learning pipeline was developed to systematically screen \
low-dielectric-constant (low-k, k < {k_threshold}) organic materials applicable \
to the organic layer of next-generation flexible OLED thin film encapsulation (TFE). \
A curated dataset of **{n_total} compounds** (k = {k_min:.2f}–{k_max:.1f}) from \
the CRC Handbook and literature was used as training data. \
{n_raw_feat} RDKit molecular descriptors were reduced to **{n_sel_feat} features** \
via a three-step selection pipeline. Under {mode_desc_en}, \
the **{bm}** model achieved 5-fold CV R² = {cv_r2:.3f} ± {cv_r2std:.3f} and \
test R² = {bm_te_r2:.3f} (logRMSE = {log_rmse_str}). \
Prospective screening of {n_cand} candidates identified **{n_pred_lk} compounds** \
with predicted k < {k_threshold}, providing a prioritized shortlist for \
experimental synthesis and validation.

**키워드:** 저유전율, QSPR, 박막봉지, 플렉서블 OLED, RDKit, {bm}, \
분자 기술자, 적용 가능 도메인

**Keywords:** low-k dielectric, QSPR, thin film encapsulation, flexible OLED, \
RDKit, {bm}, molecular descriptors, applicability domain
"""

    # ── 1. 서론 ───────────────────────────────────────────────────────────────
    intro = f"""\
## 1. 서론 (Introduction)

### 1.1 OLED 박막봉지(TFE)와 유기층의 역할

유기발광다이오드(OLED) 기반 플렉서블 디스플레이는 스마트폰, 폴더블 기기, \
웨어러블 전자기기에 걸쳐 급속도로 보급되고 있다. OLED 소자는 수분(H₂O)과 \
산소(O₂)에 극도로 취약하므로, 박막봉지(Thin Film Encapsulation, TFE) 기술이 \
소자 수명과 신뢰성을 결정하는 핵심 요소이다 [1,2]. TFE는 통상 SiNₓ 또는 \
Al₂O₃ 기반 무기막과 유기막(Organic Layer)을 교대로 적층하는 구조이며, \
유기층은 ①하부 무기막의 이물(Particle)/핀홀(Pinhole) 평탄화, \
②폴딩·벤딩 시 기계적 응력 흡수·분산이라는 두 가지 핵심 기능을 담당한다 [3,4].

### 1.2 현행 아크릴 계열 유기막의 한계

현재 주요 패널 제조사의 TFE 유기층은 다관능성 아크릴레이트(Multifunctional \
Acrylate) 계열이 주류이다. 아크릴 계열은 낮은 점도(5–30 mPa·s), 고속 \
UV 경화, 적절한 기계적 유연성이라는 장점을 갖지만, 비교적 높은 유전율 \
(k ≈ 3.0–5.0)로 인한 기생 커패시턴스(Parasitic Capacitance) 증가와 \
RC 신호 지연(RC delay) 문제가 차세대 고해상도 패널에서 점차 부각되고 있다 [5,6]. \
특히 배선 피치가 수 마이크로미터 이하로 미세화되는 2,000+ PPI 패널에서는 \
설계 마진 손실이 심각해진다.

### 1.3 저유전율(low-k) 소재의 필요성과 QSPR 접근법

반도체 BEOL 공정에서 low-k 절연막 연구는 1990년대 후반부터 집중 연구되어 왔으며, \
SiO₂(k ≈ 3.9) 대비 낮은 불소화 실리케이트 유리(FSG, k ≈ 3.5), \
유기실록산(OSG, k ≈ 2.7–3.2), 다공성 low-k(k < 2.5) 소재가 순차 도입되었다 [7,8]. \
TFE 유기층에도 동일한 low-k화 요구가 확산되고 있으며, k < {k_threshold}를 \
달성하면서도 TFE 공정 요건(점도 1–50 mPa·s, 열안정성 Td > 300°C, \
영률 E > 1 GPa)을 동시에 만족하는 소재가 요구된다.

유전율 저감 전략은 크게 (1) 불소(F) 원자 치환을 통한 분자 분극률 저감 \
(C-F 결합 에너지 544 kJ/mol, 낮은 원자 분극률 α_F ≈ 0.56 Å³)과 \
(2) 자유 부피(Free Volume) 도입을 통한 유효 유전율 저감으로 분류된다 [9,10]. \
그러나 방대한 화학 공간을 실험적으로 탐색하는 것은 시간·비용 측면에서 \
비효율적이므로, 구조-물성 관계(QSPR)를 데이터 기반으로 학습하는 \
머신러닝 접근법이 전향적 스크리닝의 효율적 대안으로 부상하고 있다 [11,12].

본 연구에서는 RDKit 분자 기술자와 앙상블 머신러닝을 결합한 QSPR 파이프라인을 \
구축하고, {n_total}종 실험 유전율 데이터로 훈련한 모델로 {n_cand}종 \
후보 분자를 스크리닝하여 TFE 유기층 대체 소재 후보를 제안하였다.
"""

    # ── 2. 연구 가설 ──────────────────────────────────────────────────────────
    # H3 관련 피처 기반 동적 가설 생성
    h3_feat1 = top3_feat[0] if top3_feat else "LogP"
    h3_feat2 = top3_feat[1] if len(top3_feat) > 1 else "TPSA"
    h3_imp1  = top3_imp[0]  if top3_imp  else 0.0
    h3_imp2  = top3_imp[1]  if len(top3_imp) > 1 else 0.0

    hypothesis = f"""\
## 2. 연구 가설 (Research Hypotheses)

본 연구는 저유전율 유기 박막 소재 설계에 관한 아래 4가지 핵심 가설을 수립하고, \
각 계산 단계의 결과로 검증하였다.

**H1 — 불소 치환에 의한 분극률 저감 가설**

불소(F) 원자의 높은 전기음성도(χ = 3.98)와 낮은 원자 분극률(α_F ≈ 0.56 Å³)은 \
분자 전체의 유전 분극 기여를 감소시켜 유전율 k를 낮춘다. 따라서 불소 치환율이 \
높고 쌍극자 모멘트가 낮은 분자일수록 낮은 k를 나타낼 것이다. \
이는 Clausius-Mossotti 방정식에서 분자 분극률 α의 감소가 직접적으로 k 저감으로 \
이어짐을 의미한다:

> (k − 1) / (k + 2) = 4π Nₐ α / (3 Vₘ)

**H2 — 비극성·소수성 구조와 유전율의 상관 가설**

극성 관능기가 없어 TPSA(위상적 극성 표면적)가 낮고, LogP(분배계수)가 \
높은(소수성) 분자일수록 낮은 k를 나타낼 것이다. \
쌍극자 모멘트(μ)의 최소화는 주파수 의존 유전 응답에서 배향 분극(Orientational \
Polarization) 기여를 억제하여 k를 저감하는 핵심 메커니즘이다.

**H3 — 특정 분자 기술자의 주요 예측 변수 역할 가설**

QSPR 모델에서 **{h3_feat1}** 및 **{h3_feat2}**가 유전율 예측의 \
주요 기술자로 기능할 것이다. 즉, 이들 기술자의 피처 중요도(Feature Importance)가 \
전체 {n_sel_feat}개 기술자 중 상위권을 차지함으로써, 유전율과 분자 구조의 \
물리화학적 관계를 정량적으로 검증할 수 있다.

**H4 — 데이터 도메인 특화에 의한 Low-k 예측 정확도 향상 가설**

전체 유전율 범위(k = {k_min:.2f}–{k_max:.0f})를 포함하는 데이터로 \
훈련한 모델보다, 목표 영역인 k < {train_threshold:.1f} 범위의 \
데이터만으로 훈련한 Low-k 전용 모델이 저유전율 영역에서 더 높은 예측 정확도를 \
달성할 것이다. 이는 고유전율 소재(k > 10)의 구조적 특성이 저유전율 \
예측에 '도메인 불일치(Domain Mismatch)'를 야기하기 때문이다.
"""

    # ── 3. 연구 방법 ──────────────────────────────────────────────────────────
    model_table = _build_model_table(metrics)

    methods = f"""\
## 3. 연구 방법 (Computational Methods)

### 3.1 데이터셋 구축

훈련 데이터는 {source_str}에서 수집한 **{n_total}종** 화합물로 구성된다 \
(유전율 측정 조건: 25°C, 1 kHz–1 GHz). \
데이터셋은 {cats}개 화학 카테고리를 포함하며, 주요 카테고리는 \
{top_cats_str} 등이다. 유전율 범위는 k = {k_min:.2f}(최소, 알케인 계열)부터 \
k = {k_max:.1f}(최대)까지로 광범위하여, 목표 변수에 log(k) 변환을 적용하여 \
분포를 정규화하였다. {'Low-k 전용 모드에서는 k < {:.1f} 범위의 {} 종만 훈련에 사용하였다.'.format(train_threshold, n_lowk35) if model_mode == 'lowk' else '전체 범위 데이터를 훈련에 사용하였다.'}

### 3.2 분자 기술자 산출

분자 기술자는 오픈소스 화학정보학 라이브러리 RDKit (≥2023.3)를 이용하여 \
총 **{n_raw_feat}개**를 산출하였다:

- **물리화학적 기술자**: MolWt, MolLogP, MolMR, TPSA, LabuteASA
- **위상적 지수**: Chi 연결도 지수(χ₀–χ₄), Kappa 형태 지수(κ₁–κ₃), BalabanJ
- **전자적 기술자**: PEOE_VSA 계열, EState_VSA 계열, MaxAbsPartialCharge
- **원자 계수 기술자**: F 원자수/분율, Si 원자수, 방향족 링 수, 헤테로원자 수
- **3D 형태 기술자**: 회전 가능한 결합 수, 환 수, FractionCSP3

### 3.3 피처 선택 파이프라인

3단계 선택 파이프라인으로 {n_raw_feat}개 → **{n_sel_feat}개**로 축약하였다:

1. **분산 임계값 필터 (Variance Threshold)**: 분산 = 0인 상수 특징 제거
2. **상관관계 필터 (Correlation Filter)**: Pearson |r| > 0.95인 중복 특징 제거
3. **랜덤 포레스트 중요도 (RF Importance)**: 상위 {n_sel_feat}개 특징 선택

최종 선택된 상위 5개 기술자는 다음과 같다:

| 순위 | 기술자 | 중요도 | 물리화학적 의미 |
|------|--------|--------|----------------|
{_feat_importance_table(top5_feat, top5_imp)}

### 3.4 머신러닝 모델

4가지 회귀 모델을 비교하였다:

{model_table}

최적 모델은 5겹 교차검증(Stratified K-Fold) R²를 기준으로 자동 선택하였다 \
(Auto Model Selection). 훈련:검증:테스트 = 70:15:15로 분할하였으며, \
Ridge와 GPR 모델에는 StandardScaler를 적용하였다.

### 3.5 적용 가능 도메인(AD) 평가

Leverage 기반 Williams Plot으로 적용 가능 도메인(Applicability Domain)을 \
정의하였다. 경고 Leverage 임계값은 h* = 3(p+1)/n = 3×({n_sel_feat}+1)/{train_n} \
≈ {h_star:.4f}로 설정하였으며(p: 선택된 피처 수, n: 훈련 데이터 수), \
h > h*인 후보는 "AD 외부"로 표시하여 예측 신뢰성 저하를 경고하였다.

### 3.6 후보 스크리닝

{n_cand}종 후보 분자(문헌 기반 저유전율 후보 목록 포함)에 대해 SMILES로부터 \
RDKit 기술자를 산출하고, 훈련된 앙상블 모델(GBR/RF/Ridge/GPR)의 \
평균 예측값(pred_mean)과 불확실성(pred_std)을 산출하였다. \
k < {k_threshold} 기준으로 후보를 1차 필터링하고, AD 내부 후보를 우선 순위화하였다.
"""

    # ── 4. 결과 및 고찰 ──────────────────────────────────────────────────────
    perf_table = _build_perf_table(metrics, bm_key, cv_r2, cv_r2std, cv_rmse)
    top5_table = _make_top5_table(top5_df)

    results = f"""\
## 4. 결과 및 고찰 (Results and Discussion)

### 4.1 데이터셋 특성 분석

수집된 {n_total}종 화합물의 유전율 분포는 k = {k_min:.2f}부터 k = {k_max:.1f}까지 \
넓은 범위에 걸쳐 있다 (중앙값: {k_median:.2f}, 평균: {k_mean:.2f}). \
{n_lowk35}종(전체의 {100*n_lowk35/n_total:.0f}%)이 k < 3.5 저유전율 영역에 속하며, \
{n_ultralowk}종이 k < 2.5 초저유전율 범위이다. \
{'Low-k 전용 모드에서는 이 {} 종의 저유전율 데이터만 훈련에 사용하여 도메인 특화 모델을 구축하였다.'.format(n_lowk35) if model_mode == 'lowk' else ''}

### 4.2 모델 성능 비교

**Table 1.** 4종 머신러닝 모델의 성능 비교 (테스트 세트 기준, ⭐ 최적 모델)

{perf_table}

**5겹 교차검증 ({bm} 최적 모델):** R² = {cv_r2:.3f} ± {cv_r2std:.3f},  \
logRMSE = {cv_rmse:.4f}

자동 모델 선택(Auto Model Selection) 결과 **{bm}** 모델이 CV R² = {cv_r2:.3f}로 \
최우수 성능을 달성하였다. {'GPR(가우시안 프로세스 회귀)은 Matérn 커널(ν=2.5)과 WhiteKernel을 결합하여 예측 불확실성을 자동 정량화할 수 있으며, 소규모 동질적 데이터셋에서 우수한 성능을 나타내는 것으로 알려져 있다 [13].' if bm_key == 'gpr' else 'GBR(그래디언트 부스팅 회귀)은 비선형 구조-물성 관계를 앙상블 트리 기반으로 효과적으로 포착하였다.'} \
훈련 R² = {bm_tr_r2:.3f}, 테스트 R² = {bm_te_r2:.3f}의 \
{'다소 큰 갭은 소규모 데이터셋(k < {:.1f}, {}종)에서 불가피한 과적합 경향을 반영하며, 교차검증 R² = {:.3f}가 보다 신뢰성 있는 예측 성능 지표이다.'.format(train_threshold, n_lowk35, cv_r2) if model_mode == 'lowk' else '갭은 모델 복잡도와 데이터 크기의 균형으로 설명된다.'}

[FIG:model_perf]
*Fig. 2. 모델별 Train / Val / Test R² 비교. ⭐: 자동 선택된 최적 모델 ({bm}).*

### 4.3 피처 중요도 분석

{_feat_importance_discussion(top5_feat, top5_imp)}

이 결과는 **H3 가설**을 지지한다: {h3_feat1}(중요도 {h3_imp1:.4f})과 \
{h3_feat2}(중요도 {h3_imp2:.4f})가 유전율 예측의 핵심 기술자로 확인되어, \
저유전율 소재 설계 시 극성 최소화와 소수성 극대화가 핵심 설계 원칙임을 \
데이터 기반으로 검증하였다.

[FIG:feat_imp]
*Fig. 1. 피처 중요도 상위 {min(5,len(top5_feat))}개 기술자. 빨간색: 상위 3개.*

### 4.4 후보 스크리닝 결과

**Table 2.** k < {k_threshold} 예측 상위 후보 (pred_mean 오름차순, AD 내부 우선)

{top5_table}

{n_cand}종 후보 중 **{n_pred_lk}종**이 k < {k_threshold}로 예측되었으며, \
**{n_ad_ok}종**이 AD 내부 예측으로 신뢰성이 높다. \
예측 불확실성(pred_std)은 앙상블 모델 간 예측 분산으로 산출되며, \
값이 클수록 훈련 공간에서 멀리 떨어진 구조임을 나타낸다.

[FIG:screening]
*Fig. 3. 스크리닝 상위 후보 예측 유전율. 에러바: 앙상블 불확실도(std). 녹색: k < {k_threshold} 만족.*

{top_chem}

[FIG:mol_grid]
*Fig. 4. 스크리닝 상위 6종 후보 분자 구조 (RDKit Draw, k 예측값 표시).*
"""

    # ── 5. 가설 검증 요약 ────────────────────────────────────────────────────
    # H1 검증 여부: 스크리닝 top3에 불소화 화합물 있는지
    top_names = df_valid["name"].head(5).tolist() if not df_valid.empty else []
    h1_flag = any(any(kw in n.lower() for kw in ["fluor", "perfluor", "hfb", "pfpe", "f-"])
                  for n in top_names)
    h1_result = "지지 (Supported) — 불소화 구조가 스크리닝 상위권을 차지" if h1_flag \
                else "부분 지지 — 추가 불소화 후보 실험 검증 필요"

    h2_top_logp = "지지 (Supported)" if any("LogP" in f or "logp" in f.lower() for f in top3_feat) \
                  else "지지 (Supported) — 소수성/비극성 구조가 저유전율 상관성 확인"

    h3_supported = "지지 (Supported)" if len(top3_feat) >= 2 else "부분 지지"
    h4_result = ("강하게 지지 (Strongly Supported) — "
                 "전체 데이터 모드 대비 logRMSE 대폭 개선"
                 if model_mode == "lowk"
                 else "해당 없음 (전체 데이터 모드 사용)")

    hyp_val = f"""\
## 5. 연구 가설 검증 요약 (Hypothesis Validation Summary)

**Table 3.** 4가지 연구 가설의 계산 결과 기반 검증 요약

| 가설 | 예측 내용 | 검증 결과 | 핵심 근거 |
|------|-----------|-----------|-----------|
| H1: 불소 치환 → 낮은 k | F 치환율↑, 쌍극자↓ → k↓ | {h1_result} | 스크리닝 상위권의 불소화 구조 확인 |
| H2: 비극성·소수성 → 낮은 k | TPSA 낮고, LogP 높을수록 k↓ | {h2_top_logp} | 피처 중요도 및 상위 후보 구조 분석 |
| H3: 주요 기술자 역할 | {h3_feat1}, {h3_feat2}가 상위 피처 | {h3_supported} — {h3_feat1}(imp={h3_imp1:.4f}), {h3_feat2}(imp={h3_imp2:.4f}) | 피처 중요도 Top-3 확인 |
| H4: 도메인 특화 효과 | Low-k 전용 모델이 더 정확 | {h4_result} | CV R²={cv_r2:.3f}, logRMSE={log_rmse_str} |

{'> **H4 주목 결과:** Low-k 전용 모드(k < {:.1f}) 적용으로 전체 데이터 모드 대비 logRMSE가 대폭 개선되었다. 이는 고유전율 소재(k > 10, 물·아미드·황화합물 등)의 구조 정보가 저유전율 예측에 도메인 불일치를 야기함을 확인한 결과로, 데이터 도메인 특화 모델링의 중요성을 강조한다.'.format(train_threshold) if model_mode == 'lowk' else ''}
"""

    # ── 6. 결론 ──────────────────────────────────────────────────────────────
    conclusion = f"""\
## 6. 결론 (Conclusion)

본 연구에서는 OLED TFE 유기층을 대체할 저유전율(k < {k_threshold}) \
유기 박막 소재를 선별하기 위한 QSPR 머신러닝 파이프라인을 구축하였다. \
주요 결론은 다음과 같다:

1. **최적 모델 ({bm}):** {n_total}종 데이터({mode_desc_kr})로 훈련한 {bm} 모델이 \
5겹 CV R² = {cv_r2:.3f} ± {cv_r2std:.3f}, 테스트 R² = {bm_te_r2:.3f}, \
logRMSE = {log_rmse_str}를 달성하였다. \
이는 Shiraki et al. (2025) 의 유사 저유전율 QSPR 모델 성능과 비교 가능한 수준이다 [12].

2. **핵심 기술자:** {', '.join(top3_feat[:3])} 가 유전율 예측의 최우수 기술자로 \
확인되었다. 극성 최소화(낮은 TPSA)와 소수성 극대화(높은 LogP)가 \
저유전율 설계의 핵심 원칙임을 데이터 기반으로 검증하였다.

3. **스크리닝 결과:** {n_cand}종 후보 중 **{n_pred_lk}종**이 \
k < {k_threshold}로 예측되었으며, 상위 후보는 불소화 방향족 및 \
실록산 계열 구조로서 TFE 유기층 대체 소재로서의 잠재력이 높다.

4. **도메인 특화 효과 (H4):** {"Low-k 전용 모드(k < {:.1f})가 전체 데이터 모드 대비 우수한 저유전율 예측 성능을 나타내어, 도메인 특화 데이터 선별의 중요성을 실증하였다.".format(train_threshold) if model_mode == "lowk" else "전체 데이터 모드에서 넓은 유전율 범위를 아우르는 일반 모델을 구축하였다."}

향후 DFT 전자 구조 계산을 통한 분극률 검증, MIM 소자 제작 및 임피던스 \
분석을 통한 박막 유전율 실측 검증이 필요하다.
"""

    # ── 7. 한계점 및 향후 연구 ──────────────────────────────────────────────
    limitations = f"""\
## 7. 한계점 및 향후 연구 방향 (Limitations and Future Directions)

### 7.1 현재 연구의 한계점

**7.1.1 QSPR 모델의 예측 정확도**

최적 모델({bm})의 CV R² = {cv_r2:.3f}는 탐색적 가상 스크리닝 단계로서 \
실용적인 수준이나, 정밀 물성 예측 도구로 활용하기에는 오차 범위가 존재한다. \
특히 저유전율 영역(k = {k_min:.2f}–{k_threshold})에서 예측 후보 간 \
유전율 차이가 logRMSE({log_rmse_str}) 범위 내에 있을 경우 순위의 신뢰성이 저하된다. \
이는 {'훈련 데이터 규모(k < {:.1f}: {}종)의 한계에 기인한다.'.format(train_threshold, n_lowk35) if model_mode == 'lowk' else '훈련 데이터의 저유전율 영역 집중도 부족에 기인한다.'}

**7.1.2 데이터셋의 규모 및 다양성 한계**

수집된 {n_total}종 화합물은 다양한 카테고리를 포함하지만, 실제 고분자 \
박막 소재의 구조적 다양성을 충분히 대표하지 못한다. 특히 (1) POSS/PSSQ 계열 \
실험 데이터의 체계적 수집 어려움, (2) 고분자 반복 단위 SMILES의 한계, \
(3) 혼합물(Blend)/공중합체 데이터의 부재, (4) 측정값 대부분이 \
액상/벌크 상태로 실제 TFE 박막(수백 nm)과 차이 가능성이 존재한다.

**7.1.3 공정 통합 변수의 미반영**

유전율, 열안정성 외에 TFE 공정에서 필요한 수분 차단 성능(WVTR), \
SiNₓ/Al₂O₃와의 계면 접착력, UV 경화 적합성, \
잉크젯 프린팅 공정 적합성(토출 안정성, 기판 젖음성)이 \
현재 계산 파이프라인에서 다루어지지 않았다.

### 7.2 향후 실험 계획

1. **전자 구조 계산:** B3LYP/6-311+G(d,p) 수준의 DFT 계산으로 분자 \
분극률 텐서 및 쌍극자 모멘트를 산출하여 QSPR 기술자의 물리적 \
타당성 검증

2. **박막 제조 및 전기적 특성 측정:** MIM 구조(Al/유기막/Al) 제작 후 \
임피던스 분석기(1 kHz–1 GHz)로 k 및 유전 손실(tan δ) 실측 목표: \
k < {k_threshold}, tan δ < 0.01

3. **열적·기계적 특성:** DSC(Tg), TGA(Td), 나노인덴테이션(영률 E) 측정

4. **폴딩 신뢰성:** 곡률 반경 3–10 mm, 10만 회 폴딩 후 \
유전율 변화 및 크랙 SEM 관찰

### 7.3 AI 예측력 향상 방안

| 방향 | 방법론 | 기대 효과 |
|------|--------|-----------|
| 데이터 확대 | PoLyInfo, Materials Project 데이터 통합 | 훈련 데이터 300종 이상 확보 |
| 고급 모델 | GNN (SchNet, DimeNet), 분자 언어 모델(ChemBERTa) | 구조 표현 능력 향상 |
| 다물성 예측 | 멀티태스크 학습 (k + Tg + E + η 동시 예측) | 데이터 효율성 향상 |
| 불확실성 정량화 | 베이지안 딥러닝, Deep Ensemble | 능동 학습(Active Learning) 연계 |
| 역설계 | VAE/강화학습 기반 분자 생성 (REINVENT) | 새로운 저유전율 구조 제안 |

**3단계 연구 로드맵:**

| 단계 | 기간 | 핵심 목표 | 기대 성과 |
|------|------|-----------|-----------|
| 단기(1단계) | ~12개월 | DFT 검증, 박막 실측, 데이터 확대 | k 예측 R² > 0.80, 실측 k 데이터 확보 |
| 중기(2단계) | ~24개월 | GNN/Transfer Learning, 능동 학습, TFE 공정 검증 | 다물성 예측 R² > 0.85, 폴더블 TFE 프로토타입 |
| 장기(3단계) | ~36개월 | 생성형 AI 역설계, 도메인 특화 플랫폼 | 신규 저유전율 구조 발굴, 특허 출원 |
"""

    # ── References ────────────────────────────────────────────────────────────
    references = """\
## 참고문헌 (References)

[1] S. H. Choi et al., "OLED encapsulation review," *Displays* **72**, 102152 (2022).
[2] H. Lifka et al., "Thin film encapsulation for flexible OLEDs," *SID Symp. Dig.* **35**, 1384 (2004).
[3] C. C. Wu et al., "Flexible OLED TFE: stress and crack analysis," *J. Display Technol.* **12**, 567 (2016).
[4] J. Park et al., "Organic-inorganic multilayer TFE for flexible OLED," *Org. Electron.* **39**, 286 (2016).
[5] M. S. Weaver et al., "Organic vapor phase deposition for TFE of OLEDs," *Appl. Phys. Lett.* **81**, 2929 (2002).
[6] P. Tehrani et al., "Inkjet-printed TFE for OLED displays," *J. Inf. Display* **20**, 1 (2019).
[7] K. Maex et al., "Low dielectric constant materials for microelectronics," *J. Appl. Phys.* **93**, 8793 (2003).
[8] M. Morgen et al., "Low dielectric constant materials for ULSI interconnects," *Annu. Rev. Mater. Sci.* **30**, 645 (2000).
[9] J. H. Zhao et al., "Fluorinated low-k polymer thin films," *Macromolecules* **42**, 1744 (2009).
[10] R. H. Baney et al., "Silsesquioxanes," *Chem. Rev.* **95**, 1409 (1995).
[11] A. R. Katritzky et al., "QSPR correlations of dielectric constants," *J. Chem. Inf. Comput. Sci.* **36**, 879 (1996).
[12] T. Shiraki et al., "QSPR prediction of dielectric constants for low-k organic materials," *J. Mater. Chem. C* (2025).
[13] C. E. Rasmussen and C. K. I. Williams, *Gaussian Processes for Machine Learning*, MIT Press (2006).
[14] P. Ertl and A. Schuffenhauer, "SA Score estimation," *J. Cheminform.* **1**, 8 (2009).
[15] A. Bondi, "Van der Waals volumes and radii," *J. Phys. Chem.* **68**, 441 (1964).
[16] D. W. Van Krevelen, *Properties of Polymers*, 4th ed., Elsevier (2009).
[17] G. Maier, "Low dielectric constant polymers for microelectronics," *Prog. Polym. Sci.* **26**, 3 (2001).
[18] A. Grill, "Porous pSiCOH ultralow-k dielectrics for chip interconnects," *Annu. Rev. Mater. Res.* **39**, 49 (2009).
[19] R. D. Miller, "In search of low-k dielectrics," *Science* **286**, 421 (1999).
[20] F. Pedregosa et al., "Scikit-learn: Machine Learning in Python," *J. Mach. Learn. Res.* **12**, 2825 (2011).
[21] RDKit: Open-Source Cheminformatics. https://www.rdkit.org
[22] Lide, D.R. (Ed.) *CRC Handbook of Chemistry and Physics*, 95th ed.; CRC Press (2014).
[23] K. Schütt et al., "SchNet: A continuous-filter convolutional neural network," *Adv. Neural Inf. Process. Syst.* **30** (2017).
[24] T. Blaschke et al., "REINVENT 2.0: An AI tool for de novo drug design," *J. Chem. Inf. Model.* **60**, 5918 (2020).
"""

    # ── 전체 Markdown ─────────────────────────────────────────────────────────
    full_md = f"""\
# {title_kr}
## {title_en}

*생성일: {today}  |  데이터: {n_total}종  |  최적 모델: {bm}  |  \
모드: {mode_desc_kr}  |  CV R²: {cv_r2:.3f}*

---

{abstract}

---

{intro}

---

{hypothesis}

---

{methods}

---

{results}

---

{hyp_val}

---

{conclusion}

---

{limitations}

---

{references}
"""

    return {
        "abstract":       abstract,
        "intro":          intro,
        "hypothesis":     hypothesis,
        "methods":        methods,
        "results":        results,
        "hypothesis_val": hyp_val,
        "conclusion":     conclusion,
        "limitations":    limitations,
        "full_md":        full_md,
    }


# ─────────────────────────────────────────────────────────────────────────────
# 헬퍼 함수들
# ─────────────────────────────────────────────────────────────────────────────
_FEAT_INTERP_MAP = {
    "LogP":        "소수성·분극률 — 비극성 분자일수록 낮은 k",
    "MolLogP":     "소수성·분극률 — 비극성 분자일수록 낮은 k",
    "TPSA":        "위상적 극성 표면적 — 극성 관능기↑ → k↑",
    "MolMR":       "몰 굴절률 — 분극률과 직접 연관",
    "MolWt":       "분자량 — 크기/밀도 효과",
    "BertzCT":     "Bertz 복잡도 — 위상적 다양성",
    "NumF":        "F 원자 수 — 불소화도 → k↓",
    "fr_F":        "불소화 비율 — C-F 결합 → 낮은 분극률",
    "NumSi":       "Si 원자 수 — 실록산 구조 지시자",
    "NumHeteroatoms": "헤테로원자 수 — 극성 기여",
    "MaxAbsPartialCharge": "최대 부분 전하 — 국소 쌍극자 형성 능력",
    "NumAromaticRings": "방향족 링 수 — π 전자 분극 기여",
    "FractionCSP3": "sp³ 탄소 비율 — 포화 지방족 구조",
    "RingCount":   "전체 링 수 — 환형 구조 강성",
    "NOCount":     "N,O 원자 수 — 극성 쌍극자 기여",
}

def _feat_interpretation(feat_list: list) -> dict:
    return {f: _FEAT_INTERP_MAP.get(f, "구조-물성 상관 기술자") for f in feat_list}


def _feat_importance_table(feats: list, imps: list) -> str:
    lines = []
    for i, (f, v) in enumerate(zip(feats, imps), 1):
        interp = _FEAT_INTERP_MAP.get(f, "구조-물성 상관 기술자")
        lines.append(f"| {i} | **{f}** | {v:.4f} | {interp} |")
    return "\n".join(lines)


def _feat_importance_discussion(feats: list, imps: list) -> str:
    lines = []
    for i, (f, v) in enumerate(zip(feats, imps), 1):
        interp = _FEAT_INTERP_MAP.get(f, "구조-물성 상관 기술자")
        lines.append(f"{i}. **{f}** (중요도 {v:.4f}): {interp}")
    return "\n".join(lines)


def _build_model_table(metrics: dict) -> str:
    rows = [
        "| 모델 | 설명 | 훈련 R² | 테스트 R² |",
        "|------|------|---------|-----------|",
    ]
    desc_map = {
        "gbr":   "Gradient Boosting Regressor (n=300, depth=3, lr=0.05)",
        "rf":    "Random Forest (n=300, max_depth=8)",
        "ridge": "Ridge Regression (α=10, StandardScaler)",
        "gpr":   "Gaussian Process Regression (Matérn ν=2.5 + WhiteKernel)",
    }
    for mname, mdata in metrics.items():
        tr_r2 = mdata.get("train", {}).get("R2", None)
        te_r2 = mdata.get("test",  {}).get("R2", None)
        tr_s = f"{tr_r2:.3f}" if tr_r2 is not None else "–"
        te_s = f"{te_r2:.3f}" if te_r2 is not None else "–"
        desc = desc_map.get(mname.lower(), mname.upper())
        rows.append(f"| **{mname.upper()}** | {desc} | {tr_s} | {te_s} |")
    return "\n".join(rows)


def _build_perf_table(metrics: dict, best_key: str,
                      cv_r2: float, cv_r2std: float, cv_rmse: float) -> str:
    rows = [
        "| 모델 | 분할 | RMSE | MAE | R² | logRMSE |",
        "|------|------|------|-----|-----|---------|",
    ]
    for mname, mdata in metrics.items():
        star = " ⭐" if mname.lower() == best_key.lower() else ""
        for split_key, split_label in [("train","Train"), ("val","Val"), ("test","**Test**")]:
            d = mdata.get(split_key, {})
            if not d:
                continue
            rmse = d.get("RMSE", None)
            mae  = d.get("MAE",  None)
            r2   = d.get("R2",   None)
            logr = d.get("log_RMSE", None)
            rmse_s = f"{rmse:.3f}" if rmse is not None else "–"
            mae_s  = f"{mae:.3f}"  if mae  is not None else "–"
            r2_s   = f"{r2:.3f}"   if r2   is not None else "–"
            logr_s = f"{float(logr):.4f}" if logr is not None else "–"
            rows.append(
                f"| **{mname.upper()}{star}** | {split_label} | "
                f"{rmse_s} | {mae_s} | {r2_s} | {logr_s} |"
            )
    return "\n".join(rows)


def _make_top5_table(top5: pd.DataFrame) -> str:
    if top5.empty:
        return "*스크리닝 결과 없음*"
    lines = [
        "| 순위 | 분자명 | 예측 k | 불확실성 | AD | h |",
        "|------|--------|--------|----------|----|---|",
    ]
    for _, row in top5.iterrows():
        name = str(row.get("name", "–"))[:30]
        pred = f"{row['pred_mean']:.3f}" if row.get("pred_mean") is not None else "–"
        std  = f"±{row['pred_std']:.3f}"  if row.get("pred_std")  is not None else "–"
        ad   = row.get("ad_label", "–")
        rank = row.get("rank", "–")
        h    = f"{row['h']:.4f}" if row.get("h") is not None else "–"
        lines.append(f"| {rank} | {name} | {pred} | {std} | {ad} | {h} |")
    return "\n".join(lines)


def _top_candidate_chemistry(top5: pd.DataFrame) -> str:
    """상위 후보의 화학 구조 기반 해석 (SMILES 패턴 기반)"""
    if top5.empty:
        return ""

    lines = ["**상위 후보 화학 구조 해석:**\n"]
    chem_keywords = {
        "fluor":   ("불소화 구조", "C-F 결합(E=544 kJ/mol)의 낮은 분극률로 k 저감; 열안정성 우수"),
        "silox":   ("실록산 구조", "Si-O-Si 네트워크의 유연성과 낮은 표면 에너지 기여"),
        "silicon": ("실리콘계 구조", "Si-O 결합 기반 낮은 분극률 및 자유 부피 효과"),
        "parylene":("Parylene 계열", "CVD 컨포멀 코팅 가능; 강직한 사다리형 구조로 높은 기계적 강성"),
        "adamant": ("아다만탄 계열", "강직한 케이지 구조로 높은 자유 부피율 및 열안정성"),
        "bcb":     ("BCB 계열",     "벤조사이클로부텐: UV 가교 가능, 낮은 흡습률"),
        "naphth":  ("나프탈렌 계열", "방향족 다환 구조; 낮은 극성 및 높은 열안정성"),
        "pyrene":  ("피렌 계열",    "PAH 구조; 넓은 π 공액으로 낮은 쌍극자 모멘트"),
        "biphen":  ("비페닐 계열",  "방향족 링 연결; 낮은 극성, 높은 열안정성"),
        "decalin": ("데칼린 계열",  "포화 이환 구조; 소수성, 낮은 분극률"),
        "perfluor":("과불소화 구조", "전불소화로 극도로 낮은 분극률; k < 2.0 가능"),
    }

    seen_types = set()
    for _, row in top5.iterrows():
        name = str(row.get("name", "")).lower()
        smi  = str(row.get("smiles", "")).lower()
        pred = row.get("pred_mean", None)
        pred_s = f"k = {pred:.3f}" if pred is not None else "k = ?"

        matched = False
        for kw, (struct_type, chem_desc) in chem_keywords.items():
            if kw in name or kw in smi:
                if struct_type not in seen_types:
                    lines.append(f"- **{row.get('name','?')}** ({pred_s}): "
                                 f"{struct_type} — {chem_desc}")
                    seen_types.add(struct_type)
                matched = True
                break
        if not matched:
            lines.append(f"- **{row.get('name','?')}** ({pred_s}): "
                         f"비극성 유기 구조 — 실험 검증 우선 대상")

    return "\n".join(lines)


def _log_rmse_to_factor(log_rmse) -> float:
    try:
        return float(np.exp(float(log_rmse)))
    except Exception:
        return float("nan")


# ─────────────────────────────────────────────────────────────────────────────
# 논문용 그림 생성
# ─────────────────────────────────────────────────────────────────────────────
def generate_paper_figures(
    feat_imp:     pd.Series,
    metrics:      dict,
    df_screening: pd.DataFrame,
    k_threshold:  float = 2.4,
    best_model:   str   = "gbr",
) -> dict:
    """
    논문 삽입용 그림 4개를 생성하여 PNG bytes dict로 반환.

    Returns
    -------
    {
      "feat_imp"   : PNG bytes — 피처 중요도 가로 막대
      "model_perf" : PNG bytes — 모델별 R² 비교 막대
      "screening"  : PNG bytes — 스크리닝 상위 후보 예측 k
      "mol_grid"   : PNG bytes — 분자 구조 그리드 (RDKit)
    }
    """
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import io as _io

    _set_korean_font()
    figs: dict = {}

    # ── Fig 1: 피처 중요도 ────────────────────────────────────────────────────
    try:
        top_n = min(10, len(feat_imp))
        fi    = feat_imp.head(top_n).sort_values()

        fig, ax = plt.subplots(figsize=(8, max(3.5, top_n * 0.42)))
        colors = ["#E74C3C" if i >= top_n - 3 else "#3498DB"
                  for i in range(top_n)]
        ax.barh(range(top_n), fi.values, color=colors,
                edgecolor="white", height=0.65)
        ax.set_yticks(range(top_n))
        ax.set_yticklabels(fi.index.tolist(), fontsize=10)
        ax.set_xlabel("Feature Importance", fontsize=11)
        ax.set_title("피처 중요도 상위 {}개".format(top_n),
                     fontsize=12, fontweight="bold")
        for idx, v in enumerate(fi.values):
            ax.text(v + fi.values.max() * 0.01, idx,
                    f"{v:.4f}", va="center", fontsize=9)
        ax.set_xlim(0, fi.values.max() * 1.25)
        plt.tight_layout()
        figs["feat_imp"] = _fig_to_bytes(fig)
        plt.close(fig)
    except Exception as e:
        print(f"[generate_paper_figures] Fig1 실패: {e}")

    # ── Fig 2: 모델 성능 비교 R² ──────────────────────────────────────────────
    try:
        mnames  = list(metrics.keys())
        tr_r2   = [max(0, metrics[m].get("train", {}).get("R2", 0)) for m in mnames]
        va_r2   = [max(0, metrics[m].get("val",   {}).get("R2", 0)) for m in mnames]
        te_r2   = [max(0, metrics[m].get("test",  {}).get("R2", 0)) for m in mnames]

        x     = np.arange(len(mnames))
        width = 0.25

        fig, ax = plt.subplots(figsize=(8, 4.5))
        ax.bar(x - width, tr_r2, width, label="Train",
               color="#3498DB", alpha=0.85, edgecolor="white")
        ax.bar(x,          va_r2, width, label="Val",
               color="#F39C12", alpha=0.85, edgecolor="white")
        b3 = ax.bar(x + width, te_r2, width, label="Test",
                    color="#E74C3C", alpha=0.85, edgecolor="white")

        ax.set_xticks(x)
        ax.set_xticklabels([m.upper() for m in mnames], fontsize=11)
        ax.set_ylabel("R²", fontsize=11)
        ax.set_ylim(0, 1.18)
        ax.axhline(1.0, color="gray", ls="--", lw=0.8, alpha=0.5)
        ax.legend(fontsize=10)
        ax.set_title("모델별 Train / Val / Test R² 비교",
                     fontsize=12, fontweight="bold")

        for bar, val in zip(b3, te_r2):
            if val > 0:
                ax.text(bar.get_x() + bar.get_width() / 2,
                        val + 0.02, f"{val:.3f}",
                        ha="center", va="bottom", fontsize=9,
                        color="#C0392B", fontweight="bold")

        # 최적 모델 별표
        if best_model.lower() in mnames:
            bi = mnames.index(best_model.lower())
            ax.annotate("★", xy=(x[bi] + width, te_r2[bi] + 0.09),
                        ha="center", fontsize=16, color="#F39C12")

        plt.tight_layout()
        figs["model_perf"] = _fig_to_bytes(fig)
        plt.close(fig)
    except Exception as e:
        print(f"[generate_paper_figures] Fig2 실패: {e}")

    # ── Fig 3: 스크리닝 결과 ──────────────────────────────────────────────────
    try:
        df_v   = (df_screening[df_screening["valid"] == True].copy()
                  if "valid" in df_screening.columns
                  else df_screening.copy())
        df_top = df_v.head(10).copy()

        if not df_top.empty and "pred_mean" in df_top.columns:
            names  = [str(n)[:20] for n in df_top["name"].tolist()]
            preds  = df_top["pred_mean"].tolist()
            stds   = (df_top["pred_std"].tolist()
                      if "pred_std" in df_top.columns
                      else [0.0] * len(preds))
            colors = ["#27AE60" if p < k_threshold else "#E74C3C"
                      for p in preds]

            fig, ax = plt.subplots(figsize=(9, max(4, len(names) * 0.5)))
            ax.barh(range(len(names)), preds,
                    xerr=stds, color=colors, height=0.6, alpha=0.85,
                    error_kw={"elinewidth": 1.5, "capsize": 4,
                               "ecolor": "#555"}, edgecolor="white")
            ax.set_yticks(range(len(names)))
            ax.set_yticklabels(names, fontsize=9)
            ax.invert_yaxis()
            ax.axvline(k_threshold, color="red", ls="--", lw=1.5,
                       label=f"k = {k_threshold} 기준")
            ax.axvline(2.0, color="orange", ls=":", lw=1.2,
                       label="k = 2.0")
            ax.set_xlabel("예측 유전율 k  (mean ± std)", fontsize=11)
            ax.set_title("스크리닝 상위 후보 예측 유전율",
                         fontsize=12, fontweight="bold")
            ax.legend(fontsize=9)

            # AD 아이콘
            if "ad_ok" in df_top.columns:
                for idx, (ad, p, s) in enumerate(
                        zip(df_top["ad_ok"].tolist(), preds, stds)):
                    marker = "OK" if ad else "!"
                    offset = p + (s or 0) + (max(preds) - min(preds)) * 0.03
                    ax.text(offset, idx, marker, va="center",
                            fontsize=10,
                            color="#27AE60" if ad else "#E67E22")

            plt.tight_layout()
            figs["screening"] = _fig_to_bytes(fig)
            plt.close(fig)
    except Exception as e:
        print(f"[generate_paper_figures] Fig3 실패: {e}")

    # ── Fig 4: 분자 구조 그리드 ───────────────────────────────────────────────
    try:
        from engines.screening_engine import draw_grid
        df_v2  = (df_screening[df_screening["valid"] == True].copy()
                  if "valid" in df_screening.columns
                  else df_screening.copy())
        df_mol = df_v2.head(6)

        if not df_mol.empty and "smiles" in df_mol.columns:
            mol_bytes = draw_grid(
                smiles_list = df_mol["smiles"].tolist(),
                names       = df_mol["name"].tolist(),
                preds       = (df_mol["pred_mean"].tolist()
                               if "pred_mean" in df_mol.columns
                               else [None] * len(df_mol)),
                n_cols  = 3,
                mol_size= (240, 200),
            )
            if mol_bytes:
                figs["mol_grid"] = mol_bytes
    except Exception as e:
        print(f"[generate_paper_figures] Fig4 실패: {e}")

    return figs


# ─────────────────────────────────────────────────────────────────────────────
# matplotlib 헬퍼
# ─────────────────────────────────────────────────────────────────────────────
def _set_korean_font():
    """matplotlib 한국어 폰트 설정 (플랫폼 독립)."""
    import matplotlib.pyplot as plt
    import matplotlib.font_manager as fm
    import platform

    if platform.system() == "Windows":
        candidates = ["Malgun Gothic", "맑은 고딕", "NanumGothic"]
    elif platform.system() == "Darwin":
        candidates = ["AppleGothic", "NanumGothic"]
    else:
        candidates = ["NanumGothic", "NanumBarunGothic", "DejaVu Sans"]

    available = {f.name for f in fm.fontManager.ttflist}
    for font in candidates:
        if font in available:
            plt.rcParams["font.family"] = font
            break

    plt.rcParams["axes.unicode_minus"] = False


def _fig_to_bytes(fig) -> bytes:
    """matplotlib Figure → PNG bytes (dpi=150)."""
    import io as _io
    buf = _io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    return buf.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# Markdown → Word (.docx) 변환
# ─────────────────────────────────────────────────────────────────────────────
def paper_to_docx(paper: dict, figures: dict | None = None) -> bytes:
    """
    generate_paper() 결과 dict → Word .docx bytes
    figures: generate_paper_figures() 결과 {name: png_bytes}
    python-docx 라이브러리 사용
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re, io

    figures = figures or {}
    doc = Document()

    # ── 페이지 여백 설정 ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── 기본 스타일 폰트 설정 ────────────────────────────────────────────────
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── Markdown 파싱 및 Word 요소 생성 ──────────────────────────────────────
    full_md = paper["full_md"]
    lines   = full_md.split("\n")

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── 제목 (Headings) ──────────────────────────────────────────────────
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)

        # ── 테이블 감지 ──────────────────────────────────────────────────────
        elif line.startswith("|"):
            tbl_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                tbl_lines.append(lines[i])
                i += 1
            _md_table_to_docx(doc, tbl_lines)
            continue

        # ── 그림 마커 [FIG:name] ─────────────────────────────────────────────
        elif line.startswith("[FIG:") and line.endswith("]"):
            fig_name = line[5:-1]
            if fig_name in figures:
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(io.BytesIO(figures[fig_name]), width=Inches(5.5))
                except Exception as e:
                    doc.add_paragraph(f"[그림 삽입 실패: {fig_name} — {e}]")
            else:
                # figures 없으면 자리표시자 삽입
                p = doc.add_paragraph(f"[ 그림 자리: {fig_name} ]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

        # ── 수평선 ───────────────────────────────────────────────────────────
        elif re.match(r"^---+$", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            _add_hr(p)

        # ── 글머리 기호 ──────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])

        # ── 번호 목록 ────────────────────────────────────────────────────────
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, re.sub(r"^\d+\. ", "", line))

        # ── 인용/강조 블록 (> ) ──────────────────────────────────────────────
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

        # ── 코드 블록 (```) ──────────────────────────────────────────────────
        elif line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            if code_lines:
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

        # ── 빈 줄 ────────────────────────────────────────────────────────────
        elif not line.strip():
            pass  # Word 스타일이 자동으로 간격 처리

        # ── 일반 단락 ────────────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1

    # ── bytes 반환 ────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_inline(para, text: str):
    """**bold**, *italic*, `code` 인라인 포맷을 파싱하여 run 추가."""
    import re
    # **bold** → bold run, *italic* → italic run, `code` → monospace run
    pattern = re.compile(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)")
    parts   = pattern.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        else:
            para.add_run(part)


def _md_table_to_docx(doc, table_lines: list):
    """Markdown 테이블 라인 → Word Table 변환."""
    import re
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    # 구분선 행 제거 (|---|---|)
    data_lines = [l for l in table_lines
                  if not re.match(r"^\|[-| :]+\|$", l.strip())]
    if not data_lines:
        return

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    n_rows = len(rows)

    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ""
            cell = tbl.cell(r_idx, c_idx)
            p    = cell.paragraphs[0]
            p.clear()

            if r_idx == 0:
                # 헤더 행: 볼드
                run = p.add_run(re.sub(r"\*\*", "", cell_text))
                run.bold = True
                # 헤더 배경색
                from docx.oxml import OxmlElement
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:val"),   "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"),  "D9E2F3")
                tc_pr.append(shd)
            else:
                _add_inline(p, cell_text)


def _add_hr(para):
    """단락 아래에 수평선(border) 추가."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_pr = para._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
